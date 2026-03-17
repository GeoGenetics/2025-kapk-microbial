#!/usr/bin/env python3
"""Remove non-useful BEAST2 analyses from docx (SNP tip-dating, blind tip-dating, UCLN)."""
from docx import Document

DOCX = '/maps/projects/caeg/people/kbd606/scratch/kapk-assm/amber/workflow_paper/methods_results_update.docx'
doc = Document(DOCX)


def delete_para(p):
    p._element.getparent().remove(p._element)


def set_para(p, text):
    p.clear()
    p.add_run(text)


# ── Methods paragraphs ────────────────────────────────────────────────────────

for p in doc.paragraphs:
    t = p.text.strip()

    # [12] Remove "Two analyses were run:" framing
    if t.startswith("Tip-dating analyses were performed with BEAST2 v2.7. Two analyses were run:"):
        set_para(p, "Tip-dating was performed with BEAST2 v2.7.")
        break

for p in doc.paragraphs:
    t = p.text.strip()

    # [13] Delete SNP tip-dating methods
    if t.startswith("SNP tip-dating:"):
        delete_para(p)
        break

for p in doc.paragraphs:
    t = p.text.strip()

    # [14] Strip "Codon tip-dating: " prefix — it's now the only analysis
    if t.startswith("Codon tip-dating: The codon alignment"):
        set_para(p, t[len("Codon tip-dating: "):])
        break

# ── Results paragraphs ────────────────────────────────────────────────────────

for p in doc.paragraphs:
    t = p.text.strip()

    # [35] Strip "Codon tip-dating: " prefix from results
    if t.startswith("Codon tip-dating: At convergence"):
        set_para(p, t[len("Codon tip-dating: "):])
        break

for p in doc.paragraphs:
    t = p.text.strip()

    # [38] Delete blind tip-dating results
    if t.startswith("Blind tip-dating (geological prior):"):
        delete_para(p)
        break

for p in doc.paragraphs:
    t = p.text.strip()

    # [42] Delete UCLN results
    if t.startswith("UCLN relaxed clock:"):
        delete_para(p)
        break

# ── Table 8: remove SNP and Blind tip-dating rows ────────────────────────────

t8 = doc.tables[8]
for row in list(t8.rows):
    first = row.cells[0].text.strip()
    if "SNP tip-dating" in first or "Blind tip-dating" in first:
        row._tr.getparent().remove(row._tr)

# ── Table 9: remove SNP clock rate and Blind dating rows ─────────────────────

t9 = doc.tables[9]
for row in list(t9.rows):
    first = row.cells[0].text.strip()
    if "SNP clock rate" in first or "Blind dating tip age" in first:
        row._tr.getparent().remove(row._tr)

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(DOCX)
print("Saved.")

# ── Verify ────────────────────────────────────────────────────────────────────

doc2 = Document(DOCX)

print("\n=== All non-empty paragraphs ===")
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t:
        print(f'[{i:2d}] {t[:150]}')

print("\n=== Table 8 ===")
for row in doc2.tables[8].rows:
    print(" | ".join(c.text.strip() for c in row.cells))

print("\n=== Table 9 (BEAST2 rows) ===")
for row in doc2.tables[9].rows:
    first = row.cells[0].text.strip()
    if any(kw in first for kw in ["BEAST2", "clock", "dating", "SNP", "Blind", "stasis", "Ancient mol"]):
        print(" | ".join(c.text.strip() for c in row.cells))
